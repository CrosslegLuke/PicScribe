import docx
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as ALIGNER


def docxmaker(iterationx, iterationy, firstpic, secondpic, thirdpic):

    rawimage = 1

    base = Document()
    margins = base.sections

    for section in margins:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    topunitnumber = iterationx
    bottomunitnumber = iterationy

    TUtext = "Unit " + topunitnumber.__str__()
    secondnametext = "Unit " + bottomunitnumber.__str__()
    BUtext = "Unit " + bottomunitnumber.__str__() + " (Homework)"

    base.add_heading(TUtext)

    top_unit = base.add_table(rows=1, cols=2)


    cell1 = top_unit.cell(0,0)
    cell2 = top_unit.cell(0,1)

    image1 = cell1.paragraphs[0]
    image2 = cell2.paragraphs[0]

    run1 = image1.add_run()
    run2 = image2.add_run()

    run1.add_picture(f'PutPicturesHere!/{firstpic}.jpg', width=Inches(3.7), height=Inches(2))
    run2.add_picture(f'PutPicturesHere!/{secondpic}.jpg', width=Inches(3.7), height=Inches(2))

    base.add_heading(BUtext)

    bot_unit = base.add_table(rows=9, cols=1)


    cell3 = bot_unit.cell(0, 0)
    image3 = cell3.paragraphs[0]
    run3 = image3.add_run()
    run3.add_picture(f"PutPicturesHere!/{thirdpic}.jpg", width=Inches(4.6), height=Inches(4))
    image3.alignment = ALIGNER.CENTER


    bot_unit.cell(1, 0).text = "Write descriptive sentences about the third picture!"
    bot_unit.cell(2,
                     0).text = "1._______________________________________________________________________________________________________________________"
    bot_unit.cell(3,
                     0).text = "2._______________________________________________________________________________________________________________________"
    bot_unit.cell(4,
                     0).text = "3._______________________________________________________________________________________________________________________"
    bot_unit.cell(5, 0).text = "Rewrite the corrected sentences."
    bot_unit.cell(6,
                     0).text = "1.________________________________________________________________________________________________________________________"
    bot_unit.cell(7,
                     0).text = "2.________________________________________________________________________________________________________________________"
    bot_unit.cell(8,
                     0).text = "3.________________________________________________________________________________________________________________________"

    outputname = TUtext + "_" + secondnametext
    outputlocation ="CreatedAssignments" + "/" + outputname + ".docx"
    base.save(outputlocation)

def piccheck(firstpic, secondpic, thirdpic):


    picchecks = [os.path.isfile(f'PutPicturesHere!/{firstpic}.jpg'),
                 os.path.isfile(f'PutPicturesHere!/{secondpic}.jpg'),
                 os.path.isfile(f'PutPicturesHere!/{thirdpic}.jpg')]


    while picchecks[0] == False:
        print(" ")
        print("Picture #" + str(firstpic) + " is missing or incorrect somehow. Make sure its name is right and file type is .jpg")
        input(".*.*.*.*.*.press Enter when issue is fixed to continue.*.*.*.*.*.")
        picchecks = [os.path.isfile(f'PutPicturesHere!/{firstpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{secondpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{thirdpic}.jpg')]
    while picchecks[1] == False:
        print(" ")
        print("Picture #" + str(secondpic) + " is missing or incorrect somehow. Make sure its name is right and file type is .jpg")
        input(".*.*.*.*.*.press Enter when issue is fixed to continue.*.*.*.*.*.")
        picchecks = [os.path.isfile(f'PutPicturesHere!/{firstpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{secondpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{thirdpic}.jpg')]
    while picchecks[2] == False:
        print(" ")
        print("Picture #" + str(thirdpic) + " is missing or incorrect somehow. Make sure its name is right and file type is .jpg")
        input(".*.*.*.*.*.press Enter when issue is fixed to continue.*.*.*.*.*.")
        picchecks = [os.path.isfile(f'PutPicturesHere!/{firstpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{secondpic}.jpg'),
                     os.path.isfile(f'PutPicturesHere!/{thirdpic}.jpg')]